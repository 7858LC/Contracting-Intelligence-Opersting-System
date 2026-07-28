"""Research module Celery tasks — Agency Intelligence Brief generation.

Phase 1 of the Executive Council research pipeline. Deliberately reads only
platform-owned, public-data-derived tables (cios/models/research.py) — never
touches any tenant's data, consistent with that module's tenant-isolation
exception being the other direction (public data made available to every
Council-member tenant, not tenant data made available platform-wide).
"""

from __future__ import annotations

import asyncio
import io
import uuid
from datetime import UTC, datetime

import boto3
import structlog
from docx import Document
from sqlalchemy import select

from cios.config import settings
from cios.tasks import celery_app

log = structlog.get_logger(__name__)


def _run(coro):  # noqa: ANN001, ANN201
    return asyncio.get_event_loop().run_until_complete(coro)


@celery_app.task(
    name="cios.tasks.research.generate_agency_intelligence_brief",
    queue="research",
    bind=True,
    max_retries=2,
    default_retry_delay=300,
)
def generate_agency_intelligence_brief(self, period_label: str | None = None) -> dict:
    """Generate one quarter's Agency Intelligence Brief covering every active
    AgencyProfile. period_label defaults to the current quarter (e.g. "2026-Q3")."""
    try:
        return _run(_async_generate_brief(period_label))
    except Exception as exc:
        log.error("agency_brief_generation_failed", error=str(exc))
        raise self.retry(exc=exc)


def _current_quarter_label(now: datetime) -> tuple[str, datetime, datetime]:
    quarter = (now.month - 1) // 3 + 1
    start_month = (quarter - 1) * 3 + 1
    period_start = datetime(now.year, start_month, 1, tzinfo=UTC)
    end_month = start_month + 2
    if end_month == 12:
        period_end = datetime(now.year, 12, 31, tzinfo=UTC)
    else:
        next_month_first = datetime(now.year, end_month + 1, 1, tzinfo=UTC)
        period_end = next_month_first
    return f"{now.year}-Q{quarter}", period_start, period_end


async def _async_generate_brief(period_label: str | None) -> dict:
    from cios.agents.base import AgentContext
    from cios.agents.research_analyst_agent import ResearchAnalystAgent
    from cios.core.database import async_session_factory
    from cios.models.research import (
        AgencyBuyingPattern,
        AgencyProfile,
        ReportStatus,
        ResearchReport,
    )
    from cios.scanners.usaspending import USASpendingScanner

    # Only current-quarter generation is supported for now; an explicit
    # period_label is accepted purely as a label override for testing/backfill,
    # the date range itself always follows the current quarter until this task
    # gains real historical-period support.
    now = datetime.now(UTC)
    label, period_start, period_end = _current_quarter_label(now)
    if period_label:
        label = period_label

    async with async_session_factory() as db:
        agencies = (
            (await db.execute(select(AgencyProfile).where(AgencyProfile.is_active == True)))  # noqa: E712
            .scalars()
            .all()
        )
        if not agencies:
            return {"error": "No active AgencyProfile rows — seed DoD/GSA/VA first"}

        sections: list[dict] = []
        errors: list[str] = []

        async with USASpendingScanner() as scanner:
            for agency in agencies:
                try:
                    aggregate = await scanner.aggregate_agency_period(
                        agency_name=agency.name,
                        period_start=period_start,
                        period_end=period_end,
                    )
                    errors.extend(aggregate.get("errors", []))

                    pattern = AgencyBuyingPattern(
                        agency_id=agency.id,
                        period_start=period_start,
                        period_end=period_end,
                        total_obligated_amount=aggregate["total_obligated_amount"],
                        award_count=aggregate["award_count"],
                        recompete_count=0,  # TODO: needs same-PIID-base award history tracking
                        top_naics_breakdown=aggregate["top_naics_breakdown"],
                        source_evidence=aggregate["source_evidence"],
                    )
                    db.add(pattern)

                    agent = ResearchAnalystAgent()
                    ctx = AgentContext(tenant_id=uuid.uuid4(), user_id=uuid.uuid4())
                    run = await agent.run(
                        ctx,
                        agency_name=agency.name,
                        period_label=label,
                        period_start=period_start.date().isoformat(),
                        period_end=period_end.date().isoformat(),
                        aggregate=aggregate,
                    )
                    sections.append({"agency": agency.name, "brief": run["result"]})
                except Exception as e:
                    errors.append(f"{agency.name}: {e}")
                    log.warning("agency_brief_section_failed", agency=agency.name, error=str(e))

        if not sections:
            return {"error": "All agency sections failed", "errors": errors}

        doc_bytes = _render_docx(label, sections)
        storage_url = _upload_to_s3(doc_bytes, label)

        report = ResearchReport(
            report_type="agency_intelligence_brief",
            period_label=label,
            status=ReportStatus.PUBLISHED.value,
            title=f"Agency Intelligence Brief — {label}",
            summary=sections[0]["brief"]["executive_summary"] if sections else None,
            storage_url=storage_url,
            published_at=now,
        )
        db.add(report)
        await db.commit()

        return {
            "report_id": str(report.id),
            "period_label": label,
            "agencies_covered": [s["agency"] for s in sections],
            "errors": errors,
        }


def _render_docx(period_label: str, sections: list[dict]) -> bytes:
    doc = Document()
    doc.add_heading(f"Agency Intelligence Brief — {period_label}", level=0)

    for section in sections:
        brief = section["brief"]
        doc.add_heading(section["agency"], level=1)
        doc.add_paragraph(brief["executive_summary"])

        doc.add_heading("Key Findings", level=2)
        for finding in brief["key_findings"]:
            doc.add_paragraph(finding, style="List Bullet")

        if brief["priority_sectors"]:
            doc.add_heading("Priority Sectors", level=2)
            for sector in brief["priority_sectors"]:
                if isinstance(sector, dict):
                    doc.add_paragraph(
                        f"{sector.get('naics', '')}: {sector.get('rationale', '')}",
                        style="List Bullet",
                    )

        doc.add_heading("Capture Implications", level=2)
        doc.add_paragraph(brief["capture_implications"])

        doc.add_paragraph(f"Data note: {brief['data_limitations']}").italic = True

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _upload_to_s3(doc_bytes: bytes, period_label: str) -> str:
    client = boto3.client(
        "s3",
        endpoint_url=settings.s3_endpoint,
        aws_access_key_id=settings.s3_access_key_id,
        aws_secret_access_key=settings.s3_secret_access_key,
        region_name=settings.s3_region,
    )
    key = f"research/agency-intelligence-brief/{period_label}.docx"
    client.put_object(
        Bucket=settings.s3_bucket_exports,
        Key=key,
        Body=doc_bytes,
        ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    return f"s3://{settings.s3_bucket_exports}/{key}"
